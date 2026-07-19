import re
import os
import json
import fitz  # PyMuPDF
import docx
from langchain_groq import ChatGroq

try:
    from backend.config import DEFAULT_LLM_MODEL
except ImportError:
    DEFAULT_LLM_MODEL = "llama-3.1-8b-instant"

def extract_text_from_pdf(file_path):
    """Extract raw text from a PDF file using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract raw text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def extract_text(file_path):
    """Wrapper to extract text based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        # Fallback to reading as text if possible
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

def classify_document_type(text):
    """
    Classifies a document's text as either 'resume' or 'jd'.
    Returns 'resume', 'jd', or 'unknown'.
    """
    if not text:
        return 'unknown'
    text_lower = text.lower()
    
    # Resume indicators
    resume_indicators = [
        "education", "academic background", "gpa", "work experience", 
        "professional summary", "projects", "certifications", "skills",
        "objective", "hobbies", "personal details", "strengths", "languages",
        "cgp", "internship", "employment history", "career summary"
    ]
    
    # Job Description indicators
    jd_indicators = [
        "job description", "responsibilities", "requirements", "qualifications",
        "we are looking for", "about the role", "key responsibilities", 
        "job summary", "duties", "reporting to", "apply now", "equal opportunity",
        "ideal candidate", "role description", "position overview", "what you will do",
        "about us", "who you are", "what we offer", "benefits", "compensation",
        "successful candidate", "join our team", "target role", "jd for"
    ]
    
    resume_score = sum(1 for word in resume_indicators if word in text_lower)
    jd_score = sum(1 for word in jd_indicators if word in text_lower)
    
    # Check for email/phone patterns (strong indicators of a candidate resume)
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    phone_pattern = r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
    
    # Exclude typical company emails like jobs@, hr@, careers@ for resume scoring
    emails = re.findall(email_pattern, text)
    has_personal_email = False
    for email in emails:
        email_lower = email.lower()
        if not any(x in email_lower for x in ["hr@", "careers@", "jobs@", "recruiting@", "apply@", "info@", "contact@"]):
            has_personal_email = True
            break
            
    if has_personal_email:
        resume_score += 3
    if re.search(phone_pattern, text):
        resume_score += 2
        
    if resume_score > jd_score:
        return 'resume'
    elif jd_score > resume_score:
        return 'jd'
    else:
        return 'unknown'

def clean_text(text):
    """Basic cleaning of text to remove redundant spacing and control characters."""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_contact_info(text, api_key=None):
    """Extract Name, Email, Phone, and LinkedIn/GitHub links from text."""
    api_key = api_key or os.getenv("GROQ_API_KEY") or os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if api_key:
        try:
            llm = ChatGroq(
                model=DEFAULT_LLM_MODEL,
                groq_api_key=api_key,
                temperature=0.0
            )
            prompt = f"""
            You are an expert contact information extractor. Analyze the following resume text and extract the candidate's contact details:
            - name (the candidate's name)
            - email (email address)
            - phone (phone number)
            - linkedin (LinkedIn profile URL)
            - github (GitHub profile URL)

            Format the output STRICTLY as a JSON object with keys: "name", "email", "phone", "linkedin", "github".
            If any field is not found, return an empty string for that field (or "Unknown Candidate" for the name if name cannot be found).
            Do not output any markdown formatting like ```json or trailing text. Output raw JSON string only.

            Resume Text:
            {text}
            """
            response = llm.invoke(prompt)
            content = response.content.strip()
            # Clean markdown code blocks if returned
            if content.startswith("```json"):
                content = content[7:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()
            
            data = json.loads(content)
            info = {
                "name": str(data.get("name", "Unknown Candidate") or "Unknown Candidate"),
                "email": str(data.get("email", "") or ""),
                "phone": str(data.get("phone", "") or ""),
                "linkedin": str(data.get("linkedin", "") or ""),
                "github": str(data.get("github", "") or "")
            }
            return info
        except Exception as e:
            print(f"Error extracting contact info via LangChain: {e}. Falling back to regex.")

    info = {
        "name": "Unknown Candidate",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": ""
    }
    
    if not text:
        return info

    # Try to extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        info["email"] = email_match.group(0)

    # Try to extract phone number (supports common formats)
    phone_match = re.search(r'(?:(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\+?\d{1,4}[-.\s]?\d{2,5}[-.\s]?\d{2,5}[-.\s]?\d{2,5})', text)
    if phone_match:
        info["phone"] = phone_match.group(0)

    # Try to extract LinkedIn and GitHub URLs
    linkedin_match = re.search(r'(https?://(www\.)?linkedin\.com/in/[\w\-]+)', text, re.IGNORECASE)
    if linkedin_match:
        info["linkedin"] = linkedin_match.group(1)

    github_match = re.search(r'(https?://(www\.)?github\.com/[\w\-]+)', text, re.IGNORECASE)
    if github_match:
        info["github"] = github_match.group(1)

    # Try to guess name
    # Usually the name is at the top. Let's look at the first few lines of the text.
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        # Avoid common section headers or emails/phones as the name
        for line in lines[:5]:
            if (len(line.split()) >= 2 and 
                len(line.split()) <= 4 and 
                "resume" not in line.lower() and 
                "curriculum" not in line.lower() and
                "page" not in line.lower() and
                "contact" not in line.lower() and
                "email" not in line.lower() and
                "phone" not in line.lower() and
                "@" not in line and
                not re.search(r'\d', line)):
                info["name"] = line
                break
                
    return info

def segment_resume(text, api_key=None):
    """
    Segment the resume text into logical parts:
    Skills, Experience, Education, Projects, Certifications, Contact.
    Returns a dictionary of sections with their text content.
    """
    api_key = api_key or os.getenv("GROQ_API_KEY") or os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if api_key:
        try:
            llm = ChatGroq(
                model=DEFAULT_LLM_MODEL,
                groq_api_key=api_key,
                temperature=0.0
            )
            prompt = f"""
            You are an expert resume parser. Analyze the following resume text and segment it into these seven sections:
            1. contact (name, email, phone, location, LinkedIn/GitHub links, portfolio)
            2. skills (technical skills, programming languages, framework, tools, soft skills)
            3. experience (work experience, job descriptions, roles, dates, achievements)
            4. education (degrees, schools, GPAs, dates)
            5. projects (project descriptions, key contributions, technologies used)
            6. certifications (certifications, licenses, courses, awards)
            7. other (any other information that does not fit into the above sections, or leave blank if none)

            Format the output STRICTLY as a JSON object with keys: "contact", "skills", "experience", "education", "projects", "certifications", "other".
            The values must be the extracted text/content corresponding to each section. Do not summarize or alter the original details, keep the extracted descriptions as close to the original text as possible.
            Do not output any markdown formatting like ```json or trailing text. Output raw JSON string only.

            Resume Text:
            {text}
            """
            response = llm.invoke(prompt)
            content = response.content.strip()
            # Clean markdown code blocks if returned
            if content.startswith("```json"):
                content = content[7:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()
            
            data = json.loads(content)
            sections = {
                "contact": str(data.get("contact", "") or ""),
                "skills": str(data.get("skills", "") or ""),
                "experience": str(data.get("experience", "") or ""),
                "education": str(data.get("education", "") or ""),
                "projects": str(data.get("projects", "") or ""),
                "certifications": str(data.get("certifications", "") or ""),
                "other": str(data.get("other", "") or "")
            }
            # Clean text in each section
            for k in sections:
                sections[k] = clean_text(sections[k])
            return sections
        except Exception as e:
            print(f"Error segmenting resume via LangChain: {e}. Falling back to rule-based parser.")

    sections = {
        "contact": "",
        "skills": "",
        "experience": "",
        "education": "",
        "projects": "",
        "certifications": "",
        "other": ""
    }
    
    if not text:
        return sections

    # Define common headers for each section (regex patterns)
    patterns = {
        "skills": r'\b(?:technical skills|key skills|skills|expertise|technologies|proficiencies|core competencies)\b',
        "experience": r'\b(?:experience|work experience|employment history|professional experience|professional background|work history)\b',
        "education": r'\b(?:education|academic background|academic qualifications|degrees|scholastic background)\b',
        "projects": r'\b(?:projects|academic projects|personal projects|portfolio|key projects)\b',
        "certifications": r'\b(?:certifications|licenses|courses|accreditations|professional development|awards)\b',
    }

    # Check if text is flattened (very few newlines)
    if text.count('\n') < 5:
        # Flat text segmentation using regex matches
        matches = []
        for sec, pat in patterns.items():
            for m in re.finditer(pat, text, re.IGNORECASE):
                matches.append((m.start(), sec, m.group(0)))
        
        matches.sort()
        
        # Filter overlapping matches
        filtered_matches = []
        last_end = -1
        for start, sec, val in matches:
            end = start + len(val)
            if start >= last_end:
                filtered_matches.append((start, sec, val))
                last_end = end
                
        if not filtered_matches:
            sections["other"] = clean_text(text)
            return sections
            
        first_start, _, _ = filtered_matches[0]
        sections["contact"] = clean_text(text[:first_start])
        
        for i in range(len(filtered_matches)):
            start, sec, val = filtered_matches[i]
            next_start = len(text) if i == len(filtered_matches) - 1 else filtered_matches[i+1][0]
            content = clean_text(text[start + len(val):next_start])
            
            if sections[sec]:
                sections[sec] += " " + content
            else:
                sections[sec] = content
                
        # Clean up sections
        for s in sections:
            sections[s] = clean_text(sections[s])
        return sections

    # Find the positions of these sections in normal newline-separated text
    lines = text.split('\n')
    section_markers = []
    
    for idx, line in enumerate(lines):
        clean_line = line.strip().lower()
        # Look for headers that are relatively short (not full sentences)
        if len(clean_line) < 40:
            for sec, pat in patterns.items():
                if re.search(pat, clean_line):
                    section_markers.append((idx, sec))
                    break

    # Sort markers by line index
    section_markers.sort()

    # If no markers were found, we will segment based on keyword matching block-by-block
    if not section_markers:
        # Fallback split
        sections["other"] = clean_text(text)
        return sections

    # Add contact info as the first implicit section before the first marker
    first_idx, first_sec = section_markers[0]
    sections["contact"] = "\n".join(lines[:first_idx])

    # Slice the lines according to markers
    for i in range(len(section_markers)):
        current_idx, current_sec = section_markers[i]
        next_idx = len(lines) if i == len(section_markers) - 1 else section_markers[i+1][0]
        
        section_content = "\n".join(lines[current_idx + 1:next_idx]).strip()
        # Append content if section is repeated (rare, but possible)
        if sections[current_sec]:
            sections[current_sec] += "\n\n" + section_content
        else:
            sections[current_sec] = section_content

    # Clean up sections
    for sec in sections:
        sections[sec] = clean_text(sections[sec])

    return sections
